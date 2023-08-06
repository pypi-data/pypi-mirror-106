import pandas as pd
import torch
import numpy as np

from docx import Document
from docx.shared import Inches

from docx import Document
from docx.shared import Inches

def make_dokument_limit(book, result, file_name, title=None, first_paragraf=None, adress = "/content/drive/MyDrive/diplom/results/", h_m=None):
  if h_m == None:
    h_m = len(book)

  document = Document()

  if title != None:
    document.add_heading(title, 0)

  if first_paragraf != None: 
    p = document.add_paragraph(first_paragraf)

  for i in range(h_m):
    ind = result[i][0, 0]
    p = document.add_paragraph(str(i+1)+"  ")
    p.add_run(book[int(ind)])
    array = result[i][1] > 1
    array = result[i][1][array]
    p.add_run(str(array)).italic = True

    
  document.save(adress + file_name)


##########################################################################################
def make_dokument(book, result, file_name, title=None, first_paragraf=None, adress = "/content/drive/MyDrive/diplom/results/", h_m=None):
  if h_m == None:
    h_m = len(book)

  document = Document()

  if title != None:
    document.add_heading(title, 0)

  if first_paragraf != None: 
    p = document.add_paragraph(first_paragraf)

  obraz_tokenized = tokenizer.tokenize(obraz)

  for i in range(h_m):
    ind = result[i][0, 0]    
    # if "кон" in book[int(ind)]:
    p = document.add_paragraph(str(i+1)+"  ")
    p.add_run(book[int(ind)])
    array = ""
    for j in np.arange(result[i][1].shape[0]):
      if result[i][1][j] < 1:
        print(result[i][1][j])
      array += obraz_tokenized[j]+", "+('%.1f' % result[i][1][j])+ "   "
    p.add_run(array).italic = True

    
  document.save(adress + file_name)
##########################################################################################
def make_result_dokument(final_array_with, final_array_without, file_name, title=None, first_paragraf=None, adress = "/content/drive/MyDrive/diplom/results/results/", h_m=None):
  if h_m == None:
    h_m = 30

  document = Document()

  if title != None:
    document.add_heading(title, 0)
  
  document.add_heading('Описание способа выбора', level=1)
  if first_paragraf != None: 
    p = document.add_paragraph(first_paragraf)

  obraz_tokenized = tokenizer.tokenize(obraz)

  document.add_heading('Лучшие тексты С ключевыми буквосочетаниями', level=1)
  i = 0
  for text in final_array_with:
    i += 1
    p = document.add_paragraph(str(i)+" "+text)
  
  document.add_heading('Лучшие тексты БЕЗ ключевых буквосочетаний', level=1)
  i = 0
  for text in final_array_without[:h_m]:
    i += 1
    p = document.add_paragraph(str(i)+" "+text)

    
  document.save(adress + file_name)